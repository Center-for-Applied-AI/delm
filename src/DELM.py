from __future__ import annotations

"""
DELM v0.2 – “Phase‑2” implementation
-----------------------------------
Major upgrades from the Phase‑1 prototype:
• Multi‑format loaders (.txt, .html/.md, .docx, .pdf*)
• Pluggable split strategies (ParagraphSplit, FixedWindowSplit, RegexSplit)
• Relevance scoring abstraction (KeywordScorer, FuzzyScorer)
• Structured extraction via `Instructor` with a Pydantic schema fallback
• Built‑in stub fallback when either OPENAI_API_KEY or Instructor is absent
*PDF uses `marker` OCR if available; else raises NotImplementedError.

The public API and method signatures remain unchanged so downstream code
continues to work.  Future phases should require **only** new strategy
classes or loader helpers – no breaking changes.
"""

import json
import os
import re
import time
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Dict, List, Sequence, Tuple, Optional

import dotenv
import pandas as pd
import yaml
from tqdm.auto import tqdm

# Required deps -------------------------------------------------------------- #
import openai  # type: ignore
import instructor  # type: ignore
from pydantic import BaseModel, Field

# Import new schema system and processing utilities
from schemas import SchemaRegistry
from processing import BatchProcessor, CostTracker, RetryHandler
from models import ExtractionVariable

try:
    from bs4 import BeautifulSoup  # type: ignore
except ImportError:  # pragma: no cover
    BeautifulSoup = None  # type: ignore

try:
    import docx  # python‑docx
except ImportError:  # pragma: no cover
    docx = None  # type: ignore

try:
    import marker  # OCR & PDF
except ImportError:  # pragma: no cover
    marker = None  # type: ignore

# --------------------------------------------------------------------------- #
# Strategy abstractions                                                       #
# --------------------------------------------------------------------------- #
class SplitStrategy(ABC):
    """Return list[str] given raw document text – override .split."""

    @abstractmethod
    def split(self, text: str) -> List[str]:
        raise NotImplementedError


class ParagraphSplit(SplitStrategy):
    # split text into paragraphs by newlines
    REGEX = re.compile(r"\r?\n\s*\r?\n")

    def split(self, text: str) -> List[str]:
        return [p.strip() for p in self.REGEX.split(text) if p.strip()]


class FixedWindowSplit(SplitStrategy):
    def __init__(self, window: int = 5, stride: int | None = None):
        self.window, self.stride = window, stride or window

    def split(self, text: str) -> List[str]:
        sentences = re.split(r"(?<=[.!?])\s+", text)
        chunks, i = [], 0
        while i < len(sentences):
            chunk = " ".join(sentences[i : i + self.window])
            chunks.append(chunk.strip())
            i += self.stride
        return [c for c in chunks if c]


class RegexSplit(SplitStrategy):
    def __init__(self, pattern: str):
        self.pattern = re.compile(pattern)

    def split(self, text: str) -> List[str]:
        return [p.strip() for p in self.pattern.split(text) if p.strip()]


class RelevanceScorer(ABC):
    @abstractmethod
    def score(self, paragraph: str) -> float:
        raise NotImplementedError


class KeywordScorer(RelevanceScorer):
    def __init__(self, keywords: Sequence[str]):
        self.keywords = [kw.lower() for kw in keywords]

    def score(self, paragraph: str) -> float:
        lowered = paragraph.lower()
        return float(any(kw in lowered for kw in self.keywords))


class FuzzyScorer(RelevanceScorer):
    def __init__(self, keywords: Sequence[str]):
        self.keywords = [kw.lower() for kw in keywords]
        try:
            from rapidfuzz import fuzz  # type: ignore
        except ImportError:
            fuzz = None  # type: ignore
        self.fuzz = fuzz

    def score(self, paragraph: str) -> float:  # 0‑1 range
        if self.fuzz is None:
            return KeywordScorer(self.keywords).score(paragraph)
        lowered = paragraph.lower()
        return max(self.fuzz.partial_ratio(lowered, kw) / 100 for kw in self.keywords)


# ExtractionVariable and ExtractionSchema moved to models.py and schemas.py


# --------------------------------------------------------------------------- #
# Helper functions for file loading                                          #
# --------------------------------------------------------------------------- #
NUMERIC_REGEX = re.compile(r"-?\d[\d,\.]*")


def load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def load_html(path: Path) -> str:
    if BeautifulSoup is None:
        raise ImportError("BeautifulSoup4 not installed but required for .html/.md loading")
    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="replace"), "html.parser")
    return soup.get_text("\n")


def load_docx(path: Path) -> str:
    if docx is None:
        raise ImportError("python-docx not installed but required for .docx loading")
    doc = docx.Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def load_pdf(path: Path) -> str:
    if marker is None:
        raise ImportError("marker (OCR) not installed – PDF loading unavailable")
    # Handle different marker API versions with type ignore
    try:
        # Try newer API
        doc = marker.Marker(str(path))  # type: ignore
        return "\n".join([p.text for p in doc.paragraphs])  # type: ignore
    except AttributeError:
        # Fallback to older API
        return "\n".join(marker.parse(str(path)).paragraphs)  # type: ignore

def load_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path)

LOADER_MAP = {
    ".txt": load_txt,
    ".md": load_html,
    ".html": load_html,
    ".htm": load_html,
    ".docx": load_docx,
    ".pdf": load_pdf,
}


# --------------------------------------------------------------------------- #
# Main class                                                                  #
# --------------------------------------------------------------------------- #
DEFAULT_KEYWORDS = ("price", "forecast", "guidance", "estimate", "expectation")


class DELM:
    """Phase‑2 extraction pipeline with pluggable strategies."""

    def __init__(
        self,
        config_path: str | Path | None = None,
        dotenv_path: str | Path | None = None,
        *,
        split_strategy: SplitStrategy | None = None,
        relevance_scorer: RelevanceScorer | None = None,
    ) -> None:
        self.CHUNK_COLUMN_NAME = "text_chunk"
        self.root_dir = Path.cwd()
        self.config: Dict[str, Any] = self._load_config(config_path)

        # Env & secrets ---------------------------------------------------- #
        if dotenv_path:
            dotenv.load_dotenv(dotenv_path)
        self.api_key: str | None = os.getenv("OPENAI_API_KEY")
        if self.api_key and openai is not None:
            openai.api_key = self.api_key

        self.model_name = self.config.get("model_name", "gpt-4o-mini")
        self.temperature = float(self.config.get("temperature", 0))
        self.max_retries = int(self.config.get("max_retries", 3))

        # Strategy objects ------------------------------------------------- #
        self.splitter: SplitStrategy = split_strategy or ParagraphSplit()
        self.scorer: RelevanceScorer = relevance_scorer or KeywordScorer(DEFAULT_KEYWORDS)

        # Schema system ---------------------------------------------------- #
        self.schema_registry = SchemaRegistry()
        self.extraction_schema = None
        
        if 'extraction' in self.config:
            self.extraction_schema = self.schema_registry.create(self.config['extraction'])

        # Enhanced processing components ------------------------------------ #
        self.batch_processor = BatchProcessor(
            batch_size=self.config.get("batch_size", 10),
            max_workers=self.config.get("max_workers", 4)
        )
        self.cost_tracker = CostTracker()
        self.retry_handler = RetryHandler(
            max_retries=self.max_retries
        )

        # Runtime artefacts ------------------------------------------------ #
        self.raw_df: pd.DataFrame | None = None
        self.processed_df: pd.DataFrame | None = None

    # ------------------------------ Public API --------------------------- #
    def prep_data_from_file(self, file_path: str | Path, target_column: str = "") -> pd.DataFrame:
        path = Path(file_path)
        loader = LOADER_MAP.get(path.suffix.lower())
        if not loader:
            raise ValueError(f"Unsupported file type: {path.suffix}")

        if path.suffix.lower() == ".csv":
            if target_column == "":
                raise ValueError("Target column is required for CSV files")
            df = load_csv(path)
        else:
            text = loader(path)
            df = pd.DataFrame({
                self.CHUNK_COLUMN_NAME: text
            })
        # self.raw_df = df
        return df

    def prep_data_from_df(self, df: pd.DataFrame, target_column: str = "", drop_target_column: bool = True) -> pd.DataFrame:
        if target_column == "":
            raise ValueError("Target column is required")

        df = df.copy()
        df[self.CHUNK_COLUMN_NAME] = df[target_column].apply(self.splitter.split)
        df = df.explode(self.CHUNK_COLUMN_NAME).reset_index(drop=True)
        df["chunk_id"] = range(len(df))
        
        if drop_target_column:
            df = df.drop(columns=[target_column])

        df["score"] = df[self.CHUNK_COLUMN_NAME].apply(self.scorer.score)
        
        return df

    def process_via_llm(self, data: pd.DataFrame, verbose: bool = False, use_batching: bool = False, use_regex_fallback: bool = False) -> pd.DataFrame:
        """Process data through LLM extraction with optional batching and regex fallback."""
        if verbose:
            print("Validating input DataFrame")
        self._validate_input_df(data)

        # TODO: add filtering for score

        paragraphs = data[self.CHUNK_COLUMN_NAME].tolist()
        
        if use_batching and self.api_key:
            results = self._process_paragraphs_batch(paragraphs, verbose, use_regex_fallback)
        else:
            results = self._process_paragraphs_sequential(paragraphs, verbose, use_regex_fallback)
        
        out = data.copy()
        out["llm_json"] = results
        if verbose:
            print(f"Processed {len(out)} rows. Saved json to `llm_json` column")
        self.processed_df = out
        return out
    
    def _process_paragraphs_sequential(self, paragraphs: List[str], verbose: bool = False, use_regex_fallback: bool = False) -> List[Dict[str, Any]]:
        """Process paragraphs sequentially."""
        results = []
        for paragraph in tqdm(paragraphs, desc="Extracting via LLM"):
            result = self._extract_from_paragraph(paragraph, verbose, use_regex_fallback)
            results.append(result)
        return results
    
    def _process_paragraphs_batch(self, paragraphs: List[str], verbose: bool = False, use_regex_fallback: bool = False) -> List[Dict[str, Any]]:
        """Process paragraphs in parallel batches."""
        return self.batch_processor.process_batch(
            paragraphs, 
            lambda p: self._extract_from_paragraph(p, verbose, use_regex_fallback)
        )
    
    def _extract_from_paragraph(self, paragraph: str, verbose: bool = False, use_regex_fallback: bool = False) -> Dict[str, Any]:
        """Extract data from a single paragraph with optional fallback."""
        if not self.api_key:
            if verbose:
                print("No API key found, falling back to regex extraction")
            return self._regex_extract(paragraph)
        
        try:
            if verbose:
                print(f"Extracting LLM for text chunk of size {len(paragraph)}")
            result = self._instructor_extract(paragraph)
            if verbose:
                print(f"LLM extraction result: {result}")
            return result
        except Exception as e:
            if verbose:
                print(f"Error processing paragraph: {e}")
            if use_regex_fallback:
                if verbose:
                    print("Falling back to regex extraction")
                return self._regex_extract(paragraph)
            else:
                if verbose:
                    print("No regex fallback enabled, returning empty result")
                # Return empty result instead of falling back to regex
                return {}

    def evaluate_output_metrics(self, data: pd.DataFrame) -> Dict[str, float]:
        self._validate_output_df(data)

        valid_ratio = float(data["llm_json"].apply(lambda x: isinstance(x, dict)).mean())
        avg_numbers = float(
            data["llm_json"].apply(lambda x: len(x.get("numbers", [])) if isinstance(x, dict) else 0).mean()
        )
        
        # Get cost summary
        cost_summary = self.cost_tracker.get_summary()
        
        return {
            "rows": float(len(data)),
            "valid_json": round(valid_ratio, 4),
            "avg_numbers": round(avg_numbers, 2),
            **cost_summary
        }
    
    def get_cost_summary(self) -> Dict[str, Any]:
        """Get detailed cost and usage summary."""
        return self.cost_tracker.get_summary()
    
    def parse_to_dataframe(self, data: pd.DataFrame, metadata: Dict[str, Any] | None = None) -> pd.DataFrame:
        """Parse LLM responses into a structured DataFrame."""
        if self.processed_df is None:
            raise ValueError("No processed data available. Run process_via_llm first.")
        
        all_results = []
        for idx, row in data.iterrows():
            response = row.get("llm_json")
            paragraph = row.get(self.CHUNK_COLUMN_NAME, "")
            
            # Add row metadata
            row_metadata = {
                'chunk_id': row.get('chunk_id', idx),
                'score': row.get('score', 0.0)
            }
            if metadata:
                row_metadata.update(metadata)
            
            # Parse response using the schema
            if self.extraction_schema and response is not None:
                # Handle both Pydantic models and dictionaries
                if hasattr(response, 'dict'):
                    # It's a Pydantic model
                    parsed_df = self.extraction_schema.parse_response(response, str(paragraph), row_metadata)
                elif isinstance(response, dict):
                    # It's a dictionary (fallback case)
                    # Try to convert to Pydantic model if possible
                    try:
                        schema_class = self.extraction_schema.create_pydantic_schema()
                        if hasattr(response, 'get') and self.extraction_schema.container_name in response:
                            # For nested schemas, we need to handle the container structure
                            parsed_df = self.extraction_schema.parse_response(response, str(paragraph), row_metadata)
                        else:
                            # For simple schemas, try to create a model instance
                            model_instance = schema_class(**response)
                            parsed_df = self.extraction_schema.parse_response(model_instance, str(paragraph), row_metadata)
                    except Exception as e:
                        print(f"Failed to parse response as Pydantic model: {e}")
                        parsed_df = pd.DataFrame()
                else:
                    parsed_df = pd.DataFrame()
                all_results.append(parsed_df)
        
        if all_results:
            return pd.concat(all_results, ignore_index=True)
        return pd.DataFrame()

    # -------------------------- Extraction helpers ----------------------- #
    def _regex_extract(self, paragraph: str) -> Dict[str, List[str]]:
        return {"numbers": NUMERIC_REGEX.findall(paragraph)}

    def _instructor_extract(self, paragraph: str) -> Dict[str, Any]:
        """Use Instructor + Pydantic schema for structured output."""
        def extract_with_schema():
            client = instructor.patch(openai.OpenAI(api_key=self.api_key))
            
            # Use configurable schema if available, otherwise create a simple default schema
            if self.extraction_schema:
                schema = self.extraction_schema.create_pydantic_schema()
                prompt = self.extraction_schema.create_prompt(paragraph)
            else:
                # Create a simple default schema for numeric extraction
                class DefaultExtractSchema(BaseModel):
                    numbers: List[str] = Field(
                        default_factory=list,
                        description="Numeric strings (keep punctuation), in order of appearance",
                    )
                schema = DefaultExtractSchema
                prompt = f"Extract all numeric strings from the following paragraph:\n\n{paragraph}"
            
            print(f"Prompt: {prompt}")

            response = client.chat.completions.create(  # type: ignore
                model=self.model_name,
                temperature=self.temperature,
                response_model=schema,
                messages=[
                    {
                        "role": "system",
                        "content": "You are a precise data‑extraction assistant.",
                        # TODO: make system prompt configurable
                    },
                    {
                        "role": "user",
                        "content": prompt,
                        # TODO: incorporate Kirill's repo prompting strategy for more comprehensive extraction
                    },
                ],
            )
            # Return the Pydantic model instance
            return response
        
        try:
            return self.retry_handler.execute_with_retry(extract_with_schema)
        except Exception as e:
            print(f"Failed to extract data from paragraph: {paragraph}.")
            print(f"Error: {e}")
            raise

    # --------------------------- Validation utils ------------------------ #
    def _validate_input_df(self, df: pd.DataFrame) -> None:
        if {self.CHUNK_COLUMN_NAME, "score"} - set(df.columns):
            raise KeyError("Input DataFrame missing required columns")

    def _validate_output_df(self, df: pd.DataFrame) -> None:
        if {self.CHUNK_COLUMN_NAME, "score", "llm_json"} - set(df.columns):
            raise KeyError("Output DataFrame missing required columns")

    # ----------------------------- Config -------------------------------- #
    @staticmethod
    def _load_config(path: str | Path | None) -> Dict[str, Any]:
        if not path:
            return {}
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(p)
        if p.suffix.lower() in {".yml", ".yaml"}:
            return yaml.safe_load(p.read_text()) or {}
        if p.suffix.lower() == ".json":
            return json.loads(p.read_text())
        raise ValueError("Config must be YAML or JSON")