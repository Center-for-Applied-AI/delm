"""
DELM Data Loaders
================
Factory pattern for loading different file formats.
"""

import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Union, Dict, Any, Callable
import pandas as pd

from delm.exceptions import DataError, FileError, DependencyError
from delm.constants import SYSTEM_FILE_NAME_COLUMN, SYSTEM_RAW_DATA_COLUMN, IGNORE_FILES

# Module-level logger
log = logging.getLogger(__name__)


# Optional dependencies
try:
    from bs4 import BeautifulSoup  # type: ignore
except ImportError:  # pragma: no cover
    BeautifulSoup = None  # type: ignore

try:
    import docx  # python‑docx
except ImportError:  # pragma: no cover
    docx = None  # type: ignore


class DataLoader(ABC):
    """Abstract base class for data loaders."""
    
    @property
    @abstractmethod
    def requires_target_column(self) -> bool:
        """Whether this loader requires a target column specification."""
        raise NotImplementedError
    
    @abstractmethod
    def load(self, path: Path) -> pd.DataFrame:
        """Load data from file and return as string or DataFrame."""
        raise NotImplementedError


class TextLoader(DataLoader):
    """Load plain text files."""
    
    @property
    def requires_target_column(self) -> bool:
        return False
    
    def load(self, path: Path) -> pd.DataFrame:
        log.debug(f"Loading text file: {path}")
        try:
            content = path.read_text(encoding="utf-8", errors="replace")
            log.debug(f"Text file loaded successfully: {path}, content length: {len(content)} characters")
            df = pd.DataFrame({SYSTEM_FILE_NAME_COLUMN: [path.name], SYSTEM_RAW_DATA_COLUMN: [content]})
            log.debug(f"Text file converted to DataFrame with {len(df)} rows")
            return df
        except Exception as e:
            log.error(f"Failed to load text file: {path}, error: {e}")
            raise


class HtmlLoader(DataLoader):
    """Load HTML/Markdown files."""
    
    @property
    def requires_target_column(self) -> bool:
        return False
    
    def load(self, path: Path) -> pd.DataFrame:
        log.debug(f"Loading HTML/Markdown file: {path}")
        if BeautifulSoup is None:
            log.error(f"BeautifulSoup4 not installed but required for .html/.md loading: {path}")
            raise DependencyError(
                "BeautifulSoup4 not installed but required for .html/.md loading",
                {"file_path": str(path), "file_type": "html/markdown"}
            )
        try:
            content = path.read_text(encoding="utf-8", errors="replace")
            log.debug(f"HTML/Markdown file read successfully: {path}, content length: {len(content)} characters")
            soup = BeautifulSoup(content, "html.parser")
            text_content = soup.get_text("\n")
            log.debug(f"HTML/Markdown parsed successfully: {path}, extracted text length: {len(text_content)} characters")
            df = pd.DataFrame({SYSTEM_FILE_NAME_COLUMN: [path.name], SYSTEM_RAW_DATA_COLUMN: [text_content]})
            log.debug(f"HTML/Markdown file converted to DataFrame with {len(df)} rows")
            return df
        except FileNotFoundError as e:
            log.error(f"HTML/Markdown file not found: {path}")
            raise FileError(f"HTML/Markdown file not found: {path}", {"file_path": str(path)}) from e
        except Exception as e:
            log.error(f"Failed to load HTML/Markdown file: {path}, error: {e}")
            raise DataError(f"Failed to load HTML/Markdown file: {path}", {"file_path": str(path)}) from e


class DocxLoader(DataLoader):
    """Load Word documents."""
    
    @property
    def requires_target_column(self) -> bool:
        return False
    
    def load(self, path: Path) -> pd.DataFrame:
        log.debug(f"Loading Word document: {path}")
        if docx is None:
            log.error(f"python-docx not installed but required for .docx loading: {path}")
            raise DependencyError(
                "python-docx not installed but required for .docx loading",
                {"file_path": str(path), "file_type": "docx"}
            )
        try:
            doc = docx.Document(str(path))
            log.debug(f"Word document opened successfully: {path}")
            text = self._extract_all_text(doc)
            log.debug(f"Word document text extracted successfully: {path}, text length: {len(text)} characters")
            df = pd.DataFrame({SYSTEM_FILE_NAME_COLUMN: [path.name], SYSTEM_RAW_DATA_COLUMN: [text]})
            log.debug(f"Word document converted to DataFrame with {len(df)} rows")
            return df
        except FileNotFoundError as e:
            log.error(f"Word document not found: {path}")
            raise FileError(f"Word document not found: {path}", {"file_path": str(path)}) from e
        except Exception as e:
            log.error(f"Failed to load Word document: {path}, error: {e}")
            raise DataError(f"Failed to load Word document: {path}", {"file_path": str(path)}) from e
    
    def _extract_all_text(self, doc) -> str:
        log.debug("Extracting text from Word document")
        text_parts = []

        # 1. Headers (for each section)
        header_count = 0
        for section in doc.sections:
            for p in section.header.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text)
                    header_count += 1
        log.debug(f"Extracted {header_count} header paragraphs")

        # 2. Main body paragraphs (includes titles/headings)
        body_count = 0
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text)
                body_count += 1
        log.debug(f"Extracted {body_count} body paragraphs")

        # 3. Tables (in order of appearance)
        table_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)
                        table_count += 1
        log.debug(f"Extracted {table_count} table cells")

        # 4. Footers (for each section)
        footer_count = 0
        for section in doc.sections:
            for p in section.footer.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text)
                    footer_count += 1
        log.debug(f"Extracted {footer_count} footer paragraphs")

        result = "\n".join(text_parts)
        log.debug(f"Word document text extraction completed: {len(text_parts)} text parts, {len(result)} total characters")
        return result


class CsvLoader(DataLoader):
    """Load CSV files."""
    
    @property
    def requires_target_column(self) -> bool:
        return True
    
    def load(self, path: Path) -> pd.DataFrame:
        log.debug(f"Loading CSV file: {path}")
        try:
            df = pd.read_csv(path)
            log.debug(f"CSV file loaded successfully: {path}, shape: {df.shape}")
            return df
        except FileNotFoundError as e:
            log.error(f"CSV file not found: {path}")
            raise FileError(f"CSV file not found: {path}", {"file_path": str(path)}) from e
        except Exception as e:
            log.error(f"Failed to load CSV file: {path}, error: {e}")
            raise DataError(f"Failed to load CSV file: {path}", {"file_path": str(path)}) from e


class ParquetLoader(DataLoader):
    """Load Parquet files."""
    @property
    def requires_target_column(self) -> bool:
        return True
    
    def load(self, path: Path) -> pd.DataFrame:
        log.debug(f"Loading Parquet file: {path}")
        try:
            df = pd.read_parquet(path)
            log.debug(f"Parquet file loaded successfully: {path}, shape: {df.shape}")
            return df
        except Exception as e:
            log.error(f"Failed to load Parquet file: {path}, error: {e}")
            raise


class FeatherLoader(DataLoader):
    """Load Feather files."""
    @property
    def requires_target_column(self) -> bool:
        return True
    
    def load(self, path: Path) -> pd.DataFrame:
        log.debug(f"Loading Feather file: {path}")
        try:
            df = pd.read_feather(path)
            log.debug(f"Feather file loaded successfully: {path}, shape: {df.shape}")
            return df
        except Exception as e:
            log.error(f"Failed to load Feather file: {path}, error: {e}")
            raise


class PdfLoader(DataLoader):
    """Load PDF files using marker OCR."""
    @property
    def requires_target_column(self) -> bool:
        return False

    def load(self, path: Path) -> pd.DataFrame:
        log.debug(f"Loading PDF file: {path}")
        try:
            import marker  # type: ignore[import]
            from marker.converters.pdf import PdfConverter  # type: ignore[import]
            from marker.models import create_model_dict  # type: ignore[import]
            from marker.output import text_from_rendered  # type: ignore[import]
            log.debug(f"Marker PDF dependencies imported successfully")
        except ImportError:
            log.error(f"marker-pdf not installed but required for .pdf loading: {path}")
            raise DependencyError(
                "marker-pdf not installed but required for .pdf loading",
                {"file_path": str(path), "file_type": "pdf"}
            )
        try:
            log.debug(f"Creating PDF converter for: {path}")
            converter = PdfConverter(artifact_dict=create_model_dict())
            log.debug(f"Converting PDF to rendered format: {path}")
            rendered = converter(str(path))
            log.debug(f"Extracting text from rendered PDF: {path}")
            text, _, _ = text_from_rendered(rendered)
            log.debug(f"PDF text extracted successfully: {path}, text length: {len(text)} characters")
            df = pd.DataFrame({SYSTEM_FILE_NAME_COLUMN: [path.name], SYSTEM_RAW_DATA_COLUMN: [text]})
            log.debug(f"PDF file converted to DataFrame with {len(df)} rows")
            return df
        except FileNotFoundError as e:
            log.error(f"PDF file not found: {path}")
            raise FileError(f"PDF file not found: {path}", {"file_path": str(path)}) from e
        except Exception as e:
            log.error(f"Failed to load PDF file: {path}, error: {e}")
            raise DataError(f"Failed to load PDF file: {path}", {"file_path": str(path)}) from e

class ExcelLoader(DataLoader):
    """Load Excel files."""
    @property
    def requires_target_column(self) -> bool:
        return True

    def load(self, path: Path) -> pd.DataFrame:
        log.debug(f"Loading Excel file: {path}")
        try:
            df = pd.read_excel(path)
            log.debug(f"Excel file loaded successfully: {path}, shape: {df.shape}")
            return df
        except Exception as e:
            log.error(f"Failed to load Excel file: {path}, error: {e}")
            raise


class DataLoaderFactory:
    """Factory for creating data loaders based on file extension."""
    
    def __init__(self):
        log.debug("Initializing DataLoaderFactory")
        self._loaders: Dict[str, DataLoader] = {
            ".txt": TextLoader(),
            ".md": HtmlLoader(),
            ".html": HtmlLoader(),
            ".htm": HtmlLoader(),
            ".docx": DocxLoader(),
            ".pdf": PdfLoader(),
            ".csv": CsvLoader(),
            ".parquet": ParquetLoader(),
            ".feather": FeatherLoader(),
            ".xlsx": ExcelLoader(),
        }
        log.debug(f"DataLoaderFactory initialized with {len(self._loaders)} loaders: {list(self._loaders.keys())}")
    
    def _get_loader(self, extension: str) -> DataLoader:
        """Get the appropriate loader for a file extension."""
        log.debug(f"Getting loader for extension: {extension}")
        loader = self._loaders.get(extension.lower())
        if loader is None:
            supported = ", ".join(self.get_supported_extensions())
            log.error(f"Unsupported file type: {extension}, supported: {supported}")
            raise DataError(
                f"Unsupported file type: {extension}",
                {
                    "file_extension": extension,
                    "supported_extensions": self.get_supported_extensions(),
                    "suggestion": f"Supported formats: {supported}"
                }
            )
        log.debug(f"Found loader for extension '{extension}': {type(loader).__name__}")
        return loader

    def get_supported_extensions(self) -> list[str]:
        """Get list of supported file extensions."""
        extensions = list(self._loaders.keys())
        log.debug(f"Supported extensions: {extensions}")
        return extensions

    def requires_target_column(self, extension: str) -> bool:
        """Check if a file extension requires a target column specification."""
        log.debug(f"Checking if extension '{extension}' requires target column")
        loader = self._loaders.get(extension.lower())
        if loader is None:
            supported = ", ".join(self.get_supported_extensions())
            log.error(f"Unsupported file type: {extension}, supported: {supported}")
            raise DataError(
                f"Unsupported file type: {extension}",
                {
                    "file_extension": extension,
                    "supported_extensions": self.get_supported_extensions(),
                    "suggestion": f"Supported formats: {supported}"
                }
            )
        requires = loader.requires_target_column
        log.debug(f"Extension '{extension}' requires target column: {requires}")
        return requires

    def _register_loader(self, extension: str, loader: DataLoader) -> None:
        """Register a new loader for a file extension."""
        log.debug(f"Registering loader for extension '{extension}': {type(loader).__name__}")
        self._loaders[extension.lower()] = loader
    
    def load_file(self, file_path: Union[str, Path]) -> pd.DataFrame:
        """Load a file using the appropriate loader."""
        path = Path(file_path)
        log.debug(f"Loading file: {path}")
        
        if not path.exists() or not path.is_file():
            log.error(f"File does not exist: {path}")
            raise FileError(f"File does not exist: {path}", {"file_path": str(path)})
        
        try:
            loader = self._get_loader(path.suffix)
            log.debug(f"Using loader for file: {path} -> {type(loader).__name__}")
        except DataError as e:
            log.error(f"Failed to get loader for file: {path}, error: {e}")
            raise FileError(f"Failed to load file: {path}", {"file_path": str(path), "error": str(e)}) from e
        
        result = loader.load(path)
        log.debug(f"File loaded successfully: {path}, result shape: {result.shape}")
        return result
    
    def load_directory(self, directory_path: Union[str, Path]) -> tuple[pd.DataFrame, str]:
        """Load a directory of files using the appropriate loader.
        
        Returns:
            tuple[pd.DataFrame, str]: A tuple containing the loaded dataframe and the extension of the loaded files.
        Raises:
            FileError: If the directory does not exist.
            FileError: If the directory contains multiple file types.
        """
        path = Path(directory_path)
        log.debug(f"Loading directory: {path}")
        
        if not path.exists():
            log.error(f"Directory does not exist: {path}")
            raise FileError(f"Directory does not exist: {path}", {"directory_path": str(path)})
        
        extensions = set()
        data = pd.DataFrame()
        file_count = 0
        
        log.debug(f"Scanning directory for files: {path}")
        for file in path.glob("**/*"):
            if file.is_file() and file.name not in IGNORE_FILES:
                log.debug(f"Loading file from directory: {file}")
                file_data = self.load_file(file)
                data = pd.concat([data, file_data], ignore_index=True)
                extensions.add(file.suffix)
                file_count += 1
            elif file.is_file():
                log.debug(f"Skipping ignored file: {file}")

        log.debug(f"Directory loading completed: {file_count} files loaded, extensions: {list(extensions)}")

        if len(extensions) != 1:
            log.error(f"Directory contains multiple file types: {path}, extensions: {list(extensions)}")
            raise FileError(
                f"Directory contains multiple file types: {path}",
                {"directory_path": str(path), "extensions found": list(extensions)}
            )
        
        extension = list(extensions)[0]
        log.debug(f"Directory loaded successfully: {path}, {file_count} files, extension: {extension}, data shape: {data.shape}")
        return data, extension

# Global factory instance
log.debug("Creating global DataLoaderFactory instance")
loader_factory = DataLoaderFactory() 