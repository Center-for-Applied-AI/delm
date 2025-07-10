"""
DELM Data Loaders
================
Factory pattern for loading different file formats.
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Union, Dict, Any, Callable
import pandas as pd

# Optional dependencies
try:
    from bs4 import BeautifulSoup  # type: ignore
except ImportError:  # pragma: no cover
    BeautifulSoup = None  # type: ignore

try:
    import docx  # python‑docx
except ImportError:  # pragma: no cover
    docx = None  # type: ignore

try:
    import marker  # OCR & PDF
except ImportError:  # pragma: no cover
    marker = None  # type: ignore


class DataLoader(ABC):
    """Abstract base class for data loaders."""
    
    @abstractmethod
    def load(self, path: Path) -> Union[str, pd.DataFrame]:
        """Load data from file and return as string or DataFrame."""
        raise NotImplementedError


class TextLoader(DataLoader):
    """Load plain text files."""
    
    def load(self, path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="replace")


class HtmlLoader(DataLoader):
    """Load HTML/Markdown files."""
    
    def load(self, path: Path) -> str:
        if BeautifulSoup is None:
            raise ImportError("BeautifulSoup4 not installed but required for .html/.md loading")
        soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="replace"), "html.parser")
        return soup.get_text("\n")


class DocxLoader(DataLoader):
    """Load Word documents."""
    
    def load(self, path: Path) -> str:
        if docx is None:
            raise ImportError("python-docx not installed but required for .docx loading")
        doc = docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)


class PdfLoader(DataLoader):
    """Load PDF files using OCR."""
    
    def load(self, path: Path) -> str:
        if marker is None:
            raise ImportError("marker (OCR) not installed – PDF loading unavailable")
        # Handle different marker API versions with type ignore
        try:
            # Try newer API
            doc = marker.Marker(str(path))  # type: ignore
            return "\n".join([p.text for p in doc.paragraphs])  # type: ignore
        except AttributeError:
            # Fallback to older API
            return "\n".join(marker.parse(str(path)).paragraphs)  # type: ignore


class CsvLoader(DataLoader):
    """Load CSV files."""
    
    def load(self, path: Path) -> pd.DataFrame:
        return pd.read_csv(path)


class DataLoaderFactory:
    """Factory for creating data loaders based on file extension."""
    
    def __init__(self):
        self._loaders: Dict[str, DataLoader] = {
            ".txt": TextLoader(),
            ".md": HtmlLoader(),
            ".html": HtmlLoader(),
            ".htm": HtmlLoader(),
            ".docx": DocxLoader(),
            ".pdf": PdfLoader(),
            ".csv": CsvLoader(),
        }
    
    def _get_loader(self, extension: str) -> DataLoader:
        """Get the appropriate loader for a file extension."""
        loader = self._loaders.get(extension.lower())
        if loader is None:
            raise ValueError(f"Unsupported file type: {extension}")
        return loader

    def get_supported_extensions(self) -> list[str]:
        """Get list of supported file extensions."""
        return list(self._loaders.keys())

    def _register_loader(self, extension: str, loader: DataLoader) -> None:
        """Register a new loader for a file extension."""
        self._loaders[extension.lower()] = loader
    
    def load_file(self, file_path: Union[str, Path]) -> Union[str, pd.DataFrame]:
        """Load a file using the appropriate loader."""
        path = Path(file_path)
        loader = self._get_loader(path.suffix)
        return loader.load(path)

# Global factory instance
loader_factory = DataLoaderFactory() 