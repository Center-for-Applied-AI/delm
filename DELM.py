from __future__ import annotations

"""
DELM v0.2 – “Phase‑2” implementation
-----------------------------------
Major upgrades from the Phase‑1 prototype:
• Multi‑format loaders (.txt, .html/.md, .docx, .pdf*)
• Pluggable split strategies (ParagraphSplit, FixedWindowSplit, RegexSplit)
• Relevance scoring abstraction (KeywordScorer, FuzzyScorer)
• Structured extraction via `Instructor` with a Pydantic schema fallback
• Built‑in stub fallback when either OPENAI_API_KEY or Instructor is absent
*PDF uses `marker` OCR if available; else raises NotImplementedError.

The public API and method signatures remain unchanged so downstream code
continues to work.  Future phases should require **only** new strategy
classes or loader helpers – no breaking changes.
"""

import json
import os
import re
import time
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Dict, List, Sequence, Tuple

import dotenv
import pandas as pd
import yaml
from tqdm.auto import tqdm

# Required deps -------------------------------------------------------------- #
import openai  # type: ignore
import instructor  # type: ignore
from pydantic import BaseModel, Field

try:
    from bs4 import BeautifulSoup  # type: ignore
except ImportError:  # pragma: no cover
    BeautifulSoup = None  # type: ignore

try:
    import docx  # python‑docx
except ImportError:  # pragma: no cover
    docx = None  # type: ignore

try:
    import marker  # OCR & PDF
except ImportError:  # pragma: no cover
    marker = None  # type: ignore

# --------------------------------------------------------------------------- #
# Strategy abstractions                                                       #
# --------------------------------------------------------------------------- #
class SplitStrategy(ABC):
    """Return list[str] given raw document text – override .split."""

    @abstractmethod
    def split(self, text: str) -> List[str]:
        raise NotImplementedError


class ParagraphSplit(SplitStrategy):
    # split text into paragraphs by newlines
    REGEX = re.compile(r"\r?\n\s*\r?\n")

    def split(self, text: str) -> List[str]:
        return [p.strip() for p in self.REGEX.split(text) if p.strip()]


class FixedWindowSplit(SplitStrategy):
    def __init__(self, window: int = 5, stride: int | None = None):
        self.window, self.stride = window, stride or window

    def split(self, text: str) -> List[str]:
        sentences = re.split(r"(?<=[.!?])\s+", text)
        chunks, i = [], 0
        while i < len(sentences):
            chunk = " ".join(sentences[i : i + self.window])
            chunks.append(chunk.strip())
            i += self.stride
        return [c for c in chunks if c]


class RegexSplit(SplitStrategy):
    def __init__(self, pattern: str):
        self.pattern = re.compile(pattern)

    def split(self, text: str) -> List[str]:
        return [p.strip() for p in self.pattern.split(text) if p.strip()]


class RelevanceScorer(ABC):
    @abstractmethod
    def score(self, paragraph: str) -> float:
        raise NotImplementedError


class KeywordScorer(RelevanceScorer):
    def __init__(self, keywords: Sequence[str]):
        self.keywords = [kw.lower() for kw in keywords]

    def score(self, paragraph: str) -> float:
        lowered = paragraph.lower()
        return float(any(kw in lowered for kw in self.keywords))


class FuzzyScorer(RelevanceScorer):
    def __init__(self, keywords: Sequence[str]):
        self.keywords = [kw.lower() for kw in keywords]
        try:
            from rapidfuzz import fuzz  # type: ignore
        except ImportError:
            fuzz = None  # type: ignore
        self.fuzz = fuzz

    def score(self, paragraph: str) -> float:  # 0‑1 range
        if self.fuzz is None:
            return KeywordScorer(self.keywords).score(paragraph)
        lowered = paragraph.lower()
        return max(self.fuzz.partial_ratio(lowered, kw) / 100 for kw in self.keywords)


# --------------------------------------------------------------------------- #
# Configurable extraction schema system                                       #
# --------------------------------------------------------------------------- #
class ExtractionVariable:
    """Defines a variable to extract from text."""
    
    def __init__(self, name: str, description: str, data_type: str, required: bool = False, allowed_values: List[str] | None = None):
        self.name = name
        self.description = description
        self.data_type = data_type
        self.required = required
        self.allowed_values = allowed_values or []
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'ExtractionVariable':
        return cls(
            name=data['name'],
            description=data['description'],
            data_type=data['data_type'],
            required=data.get('required', False),
            allowed_values=data.get('allowed_values', None)
        )


class ExtractionSchema:
    """Configurable extraction schema with variables and prompt."""
    
    def __init__(self, variables: List[ExtractionVariable], prompt_template: str):
        self.variables = variables
        self.prompt_template = prompt_template
    
    @classmethod
    def from_config(cls, config: Dict[str, Any]) -> 'ExtractionSchema':
        variables = [ExtractionVariable.from_dict(v) for v in config.get('variables', [])]
        prompt_template = config.get('prompt_template', 'Extract the following information from the text: {text}')
        return cls(variables, prompt_template)
    
    def create_pydantic_schema(self):
        """Dynamically create Pydantic schema from variables."""
        if instructor is None:
            return None
        
        # Create a proper dynamic class with proper annotations
        class DynamicExtractSchema(BaseModel):
            pass
        
        # Add fields dynamically with proper annotations
        for var in self.variables:
            if var.data_type == 'string':
                field_type = List[str]
            elif var.data_type == 'number':
                field_type = List[float]
            elif var.data_type == 'integer':
                field_type = List[int]
            elif var.data_type == 'date':
                field_type = List[str]  # Could be enhanced with datetime
            else:
                field_type = List[str]
            
            # Add validation for allowed values if specified
            field_kwargs = {
                'default_factory': list,
                'description': var.description
            }
            
            if var.allowed_values:
                field_kwargs['description'] += f" (must be one of: {', '.join(var.allowed_values)})"
                # For now, keep it as List[str] to avoid Literal complexity
                # The allowed values are enforced through the description
            
            # Use setattr to properly set the field with annotation
            setattr(DynamicExtractSchema, var.name, Field(**field_kwargs))
            # Set the annotation using __annotations__
            if not hasattr(DynamicExtractSchema, '__annotations__'):
                DynamicExtractSchema.__annotations__ = {}
            DynamicExtractSchema.__annotations__[var.name] = field_type
        
        return DynamicExtractSchema
    
    def create_prompt(self, text: str) -> str:
        """Create extraction prompt from template and variables."""
        variable_descriptions = []
        for var in self.variables:
            desc = f"- {var.name}: {var.description} ({var.data_type})"
            if var.required:
                desc += " [REQUIRED]"
            
            # Add allowed values to description so the model knows what to extract
            if var.allowed_values:
                allowed_list = ", ".join([f'"{v}"' for v in var.allowed_values])
                desc += f" (allowed values: {allowed_list})"
            
            variable_descriptions.append(desc)
        
        variables_text = "\n".join(variable_descriptions)
        return self.prompt_template.format(
            text=text,
            variables=variables_text
        )


# --------------------------------------------------------------------------- #
# Helper functions for file loading                                          #
# --------------------------------------------------------------------------- #
NUMERIC_REGEX = re.compile(r"-?\d[\d,\.]*")


def load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def load_html(path: Path) -> str:
    if BeautifulSoup is None:
        raise ImportError("BeautifulSoup4 not installed but required for .html/.md loading")
    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="replace"), "html.parser")
    return soup.get_text("\n")


def load_docx(path: Path) -> str:
    if docx is None:
        raise ImportError("python-docx not installed but required for .docx loading")
    doc = docx.Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


def load_pdf(path: Path) -> str:
    if marker is None:
        raise ImportError("marker (OCR) not installed – PDF loading unavailable")
    return "\n".join(marker.parse(path).paragraphs)

def load_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path)

LOADER_MAP = {
    ".txt": load_txt,
    ".md": load_html,
    ".html": load_html,
    ".htm": load_html,
    ".docx": load_docx,
    ".pdf": load_pdf,
}


# --------------------------------------------------------------------------- #
# Main class                                                                  #
# --------------------------------------------------------------------------- #
DEFAULT_KEYWORDS = ("price", "forecast", "guidance", "estimate", "expectation")


class DELM:
    """Phase‑2 extraction pipeline with pluggable strategies."""

    def __init__(
        self,
        config_path: str | Path | None = None,
        dotenv_path: str | Path | None = None,
        *,
        split_strategy: SplitStrategy | None = None,
        relevance_scorer: RelevanceScorer | None = None,
    ) -> None:
        self.CHUNK_COLUMN_NAME = "text_chunk"
        self.root_dir = Path.cwd()
        self.config: Dict[str, Any] = self._load_config(config_path)

        # Env & secrets ---------------------------------------------------- #
        if dotenv_path:
            dotenv.load_dotenv(dotenv_path)
        self.api_key: str | None = os.getenv("OPENAI_API_KEY")
        if self.api_key and openai is not None:
            openai.api_key = self.api_key

        self.model_name = self.config.get("model_name", "gpt-4o-mini")
        self.temperature = float(self.config.get("temperature", 0))
        self.max_retries = int(self.config.get("max_retries", 3))

        # Strategy objects ------------------------------------------------- #
        self.splitter: SplitStrategy = split_strategy or ParagraphSplit()
        self.scorer: RelevanceScorer = relevance_scorer or KeywordScorer(DEFAULT_KEYWORDS)

        # Extraction schema ------------------------------------------------- #
        self.extraction_schema: ExtractionSchema | None = None
        if 'extraction' in self.config:
            self.extraction_schema = ExtractionSchema.from_config(self.config['extraction'])

        # Runtime artefacts ------------------------------------------------ #
        self.raw_df: pd.DataFrame | None = None
        self.processed_df: pd.DataFrame | None = None

    # ------------------------------ Public API --------------------------- #
    def prep_data_from_file(self, file_path: str | Path, target_column: str = "") -> pd.DataFrame:
        path = Path(file_path)
        loader = LOADER_MAP.get(path.suffix.lower())
        if not loader:
            raise ValueError(f"Unsupported file type: {path.suffix}")

        if path.suffix.lower() == ".csv":
            if target_column == "":
                raise ValueError("Target column is required for CSV files")
            df = load_csv(path)
        else:
            text = loader(path)
            df = pd.DataFrame({
                self.CHUNK_COLUMN_NAME: text
            })
        # self.raw_df = df
        return df

    def prep_data_from_df(self, df: pd.DataFrame, target_column: str = "", drop_target_column: bool = True) -> pd.DataFrame:
        if target_column == "":
            raise ValueError("Target column is required")

        df = df.copy()
        df[self.CHUNK_COLUMN_NAME] = df[target_column].apply(self.splitter.split)
        df = df.explode(self.CHUNK_COLUMN_NAME).reset_index(drop=True)
        df["chunk_id"] = range(len(df))
        
        if drop_target_column:
            df = df.drop(columns=[target_column])

        df["score"] = df[self.CHUNK_COLUMN_NAME].apply(self.scorer.score)
        
        return df

    def process_via_llm(self, data: pd.DataFrame, verbose: bool = False) -> pd.DataFrame:
        if verbose:
            print("Validating input DataFrame")
        self._validate_input_df(data)

        # TODO: add filtering for score

        results: List[Dict[str, Any]] = []
        for paragraph in tqdm(data[self.CHUNK_COLUMN_NAME].tolist(), desc="Extracting via LLM"):
            if self.api_key:
                if verbose:
                    print(f"Extracting LLM for text chunk of size {len(paragraph)}")
                json_obj = self._instructor_extract(paragraph)
                if verbose:
                    print(f"LLM extraction result: {json_obj}")
            else:
                print("No API key found, falling back to regex extraction")
                json_obj = self._regex_extract(paragraph)
            results.append(json_obj)
        out = data.copy()
        out["llm_json"] = results
        if verbose:
            print(f"Processed {len(out)} rows. Saved json to `llm_json` column")
        self.processed_df = out
        return out

    def evaluate_output_metrics(self, data: pd.DataFrame) -> Dict[str, float]:
        self._validate_output_df(data)

        valid_ratio = float(data["llm_json"].apply(lambda x: isinstance(x, dict)).mean())
        avg_numbers = float(
            data["llm_json"].apply(lambda x: len(x.get("numbers", [])) if isinstance(x, dict) else 0).mean()
        )
        return {
            "rows": float(len(data)),
            "valid_json": round(valid_ratio, 4),
            "avg_numbers": round(avg_numbers, 2),
        }

    # -------------------------- Extraction helpers ----------------------- #
    def _regex_extract(self, paragraph: str) -> Dict[str, List[str]]:
        return {"numbers": NUMERIC_REGEX.findall(paragraph)}

    def _instructor_extract(self, paragraph: str) -> Dict[str, Any]:
        """Use Instructor + Pydantic schema for structured output."""
        attempt = 0
        while attempt < self.max_retries:
            attempt += 1
            try:
                client = instructor.patch(openai.OpenAI(api_key=self.api_key))
                
                # Use configurable schema if available, otherwise create a simple default schema
                if self.extraction_schema:
                    schema = self.extraction_schema.create_pydantic_schema()
                    prompt = self.extraction_schema.create_prompt(paragraph)
                else:
                    # Create a simple default schema for numeric extraction
                    class DefaultExtractSchema(BaseModel):
                        numbers: List[str] = Field(
                            default_factory=list,
                            description="Numeric strings (keep punctuation), in order of appearance",
                        )
                    schema = DefaultExtractSchema
                    prompt = f"Extract all numeric strings from the following paragraph:\n\n{paragraph}"
                
                print(f"Prompt: {prompt}")

                response = client.chat.completions.create(  # type: ignore
                    model=self.model_name,
                    temperature=self.temperature,
                    response_model=schema,
                    messages=[
                        {
                            "role": "system",
                            "content": "You are a precise data‑extraction assistant.",
                            # TODO: make system prompt configurable
                        },
                        {
                            "role": "user",
                            "content": prompt,
                            # TODO: incorporate Kirill's repo prompting strategy for more comprehensive extraction
                        },
                    ],
                )
                return response.dict()
            except Exception as e:
                if attempt >= self.max_retries:
                    print(f"Failed to extract data from paragraph: {paragraph}. Falling back to regex extraction.")
                    # TODO: Give details on why exactly it failed
                    # Goal: Print which variables were not extracted
                    print(f"Error: {e}")
                    return self._regex_extract(paragraph)
                # Exponential backoff for api rate limits
                time.sleep(2 ** attempt)
        return self._regex_extract(paragraph)

    # --------------------------- Validation utils ------------------------ #
    def _validate_input_df(self, df: pd.DataFrame) -> None:
        if {self.CHUNK_COLUMN_NAME, "score"} - set(df.columns):
            raise KeyError("Input DataFrame missing required columns")

    def _validate_output_df(self, df: pd.DataFrame) -> None:
        if {self.CHUNK_COLUMN_NAME, "score", "llm_json"} - set(df.columns):
            raise KeyError("Output DataFrame missing required columns")

    # ----------------------------- Config -------------------------------- #
    @staticmethod
    def _load_config(path: str | Path | None) -> Dict[str, Any]:
        if not path:
            return {}
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(p)
        if p.suffix.lower() in {".yml", ".yaml"}:
            return yaml.safe_load(p.read_text()) or {}
        if p.suffix.lower() == ".json":
            return json.loads(p.read_text())
        raise ValueError("Config must be YAML or JSON")


# # Test Code
# TEST_KEYWORDS = (
# "price",
# "prices",
# "oil",
# "gas",
# "expect",
# "barrel",
# "ton",
# "used",
# "expectations",
# "using"
# )

# DELM_CONFIG_PATH = Path("example.delm_config.yaml")
# DOTENV_PATH = Path(".env")
# TEST_FILE_PATH = Path("data/input/input2.csv")
# report_text_df = pd.read_csv(TEST_FILE_PATH).iloc[:10]
# report_text_df = report_text_df.drop(columns=["Unnamed: 0"])
# # The date is given in an inconsistent format, so it is cropped at 10 characters.
# date_clean = pd.to_datetime(report_text_df["date"].astype(str).apply(lambda x: x[:10]))
# report_text_df["date"] = date_clean
# report_text_df = report_text_df[["report", "date", "title", "subtitle", "firm_name", "text"]]
# print(report_text_df.head())
# print(report_text_df.info())
# print(report_text_df.columns)
# # print(len(report_text_df.iloc[0]["text"].split(" ")))

# # Assuming DELM is the class defined above, instantiate and use it to process input_df
# delm = DELM(
#     config_path=DELM_CONFIG_PATH, 
#     dotenv_path=DOTENV_PATH, 
#     split_strategy=ParagraphSplit(),
#     relevance_scorer=KeywordScorer(TEST_KEYWORDS)
# )
# output_df = delm.prep_data_from_df(report_text_df, "text")

# # Histogram of score
# import matplotlib.pyplot as plt
# plt.hist(output_df["score"])
# plt.show()

# llm_output_df = delm.process_via_llm(output_df.iloc[:2], verbose=True)

# print(output_df.head())
# print(output_df.info())
# print(output_df.columns)
# print(output_df.iloc[0]["text_chunk"])
# # print(len(output_df.iloc[0]["text_chunk"].split(" ")))


# # look at expected output
# test_output_df = pd.read_excel("data/output/output.xlsx").iloc[:100]
# print(test_output_df.head())
# print(test_output_df.info())
# print(test_output_df.columns)
# print(test_output_df.iloc[0])



